"""
인용발명 라우터 모듈

RAG 기반 인용발명 텍스트/파일은 필요할 때만 사용하고, 이전 사용횟수는
생성된 후 파일 DB에서 다시 조회해 original_text를 기준으로 비교한다.
"""
from __future__ import annotations

import asyncio
import gc
import hashlib
import json
import logging
import re
import shutil
import sqlite3
import time
import traceback
import uuid
from pathlib import Path
from typing import AsyncGenerator, List, Optional

import aiofiles
from fastapi import APIRouter, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from sse_starlette.sse import EventSourceResponse

from backend.models.schemas import (
    BatchDependentRequest,
    ChatRequest,
    ExtractedDocument,
    ManualClaimRequest,
    MissingPriorArtSearchRequest,
    ParsedClaim,
)
from backend.paths import CASES_DIR, REPORTS_DIR, UPLOADS_DIR
from backend.routers.settings import _load as load_settings
from backend.services import pdf_extractor
from backend.services.ai_engine import call_ai
from backend.services.ai_engine import kill_active_cli_procs
from backend.services.citation_chain import (
    CITATION_CHAIN_POLICY_VERSION,
    build_citation_chain_from_comparisons,
    get_claim_chain_info,
)
from backend.services.citation_extractor import (
    CompareFailed,
    analyze_claim_elements,
    analyze_claim_elements_for_docs,
    analyze_claim_elements_hybrid,
    get_cached_doc_indices,
    get_matches_from_cache,
    reset_incompatible_comparison_caches,
    select_candidate_doc_indices_for_elements,
    verify_quotes,
)
from backend.services.reference_store import (
    save_case_artifacts_sqlite,
    save_reference_entries_sqlite,
)
from backend.services.report_generator import (
    _dedupe_phase1_sections,
    _strip_agent_tool_calls,
    detect_category_same_claims,
    enforce_phase1_judgment_headers,
    enhance_claim_parsing_with_llm,
    enhance_purpose_effects_with_llm,
    find_unselected_reference_mentions,
    generate_category_same_report,
    generate_dependent_report,
    generate_dependent_reports_batch,
    generate_independent_phase1_streaming,
    build_rejected_inventions_section,
    format_rejection_basis_header,
    parse_manual_claim_locally,
    polish_phase1_summary_text,
    sanitize_report_status_icons,
)
from backend.services.prior_art_search import run_adaptive_prior_art_search

router = APIRouter()
logger = logging.getLogger(__name__)

DOC_CACHE_DIR = UPLOADS_DIR / "_doc_cache"
for _dir in (UPLOADS_DIR, REPORTS_DIR, CASES_DIR, DOC_CACHE_DIR):
    _dir.mkdir(exist_ok=True)

_LEGACY_FINAL_BOUNDARY_RE = re.compile(r"(?im)^\s*#\s*\[Phase\s*2\]")
_REJECTION_BASIS_HEADER_RE = re.compile(
    r"^\s*\[[^\]\r\n]*(?:신규성|진보성)[^\]\r\n]*\]"
)
_BATCH_SPLIT_RE = re.compile(r"(?m)^\s*(?:#{1,6}\s*)?===\s*청구항\s*(\d+)\s*===\s*$")


def _ev(event: str, data: str | dict) -> dict:
    return {"event": event, "data": data if isinstance(data, str) else json.dumps(data, ensure_ascii=False)}


def _elapsed(start: float) -> str:
    return f"{time.perf_counter() - start:.1f}s"


async def _await_with_log_heartbeat(coro, emit_log, *, label: str, interval: float = 15.0):
    """Keep the UI warm while a long-running awaitable is in progress."""
    task = asyncio.create_task(coro)
    started = time.perf_counter()
    try:
        while not task.done():
            try:
                return await asyncio.wait_for(task, timeout=interval)
            except asyncio.TimeoutError:
                await emit_log(f"{label} 진행 중... ({_elapsed(started)})")
        return await task
    finally:
        if not task.done():
            task.cancel()


async def _await_with_batch_status_heartbeat(
    coro,
    *,
    job_id: str,
    claim_numbers: list[int],
    started_at: str,
    stage: str,
    message_builder,
    reports_ready_getter,
    interval: float = 15.0,
):
    """Refresh polled batch status while a long-running awaitable is in progress."""
    task = asyncio.create_task(coro)
    started = time.perf_counter()
    try:
        while not task.done():
            try:
                return await asyncio.wait_for(asyncio.shield(task), timeout=interval)
            except asyncio.TimeoutError:
                _update_dependent_batch_status(
                    job_id,
                    state="running",
                    claim_numbers=claim_numbers,
                    stage=stage,
                    message=message_builder(_elapsed(started)),
                    started_at=started_at,
                    reports_ready=reports_ready_getter(),
                )
        return await task
    finally:
        if not task.done():
            task.cancel()


def _find_legacy_final_boundary(body: str) -> int:
    match = _LEGACY_FINAL_BOUNDARY_RE.search(body or "")
    return match.start() if match else -1


def _phase1_only_report(body: str) -> str:
    idx = _find_legacy_final_boundary(body)
    if idx < 0:
        return body
    return body[:idx].rstrip()


def _matches_directly_cover_single_reference(matches, reference_idx: int) -> bool:
    return bool(matches) and all(
        match.cited_invention_index == reference_idx
        and match.judgment in {"동일", "실질적 동일"}
        and match.directness in {"", "direct"}
        and not match.missing_limitations
        and bool(match.quote)
        for match in matches
    )


def _is_novelty_basis(chain_info, matches, total: list[int]) -> bool:
    """Require the parent claim and the dependent limitation to share one novelty reference."""
    if not total or len(total) != 1:
        return False
    info = chain_info or {}
    family_selection = info.get("family_selection") or {}
    analysis_track = info.get("analysis_track") or family_selection.get("analysis_track")
    is_dependent = info.get("parent") is not None
    if is_dependent:
        return (
            analysis_track == "novelty_single_reference"
            and _matches_directly_cover_single_reference(matches, total[0])
        )
    if analysis_track == "novelty_single_reference":
        return True
    return _matches_directly_cover_single_reference(matches, total[0])


def _rejection_basis_header(chain_info, matches) -> str:
    mapping = (chain_info or {}).get("doc_name_mapping", {})
    total = list((chain_info or {}).get("total") or [])
    if not total:
        total = sorted({m.cited_invention_index for m in matches}) or [0]
    inv1_idx = total[0]
    inv2_idx = total[1] if len(total) > 1 else None
    inv3_idx = total[2] if len(total) > 2 else None
    inv1_name = mapping.get(str(inv1_idx), f"인용발명 {inv1_idx + 1}")
    inv2_name = mapping.get(str(inv2_idx), f"인용발명 {inv2_idx + 1}") if inv2_idx is not None else ""
    inv3_name = mapping.get(str(inv3_idx), f"인용발명 {inv3_idx + 1}") if inv3_idx is not None else ""
    is_novelty = _is_novelty_basis(chain_info, matches, total)
    return format_rejection_basis_header(
        inv1_name,
        inv2_name,
        inv3_name,
        is_combo=len(total) > 1,
        chain_info=chain_info,
        is_novelty=is_novelty,
    )


def _prepend_rejection_basis_header(report_md: str, chain_info, matches) -> str:
    report_md = (report_md or "").strip()
    report_md = re.sub(
        r"(?im)^\s*#{1,6}\s*(?:[^\w\r\n]*)?\s*종합\s*분석\s*요약\s*$",
        "[종합분석요약]",
        report_md,
    )
    if _REJECTION_BASIS_HEADER_RE.match(report_md):
        lines = report_md.splitlines()
        if len(lines) > 1 and lines[1].strip() == "[구성대비]":
            return report_md
        if len(lines) > 2 and not lines[1].strip() and lines[2].strip() == "[구성대비]":
            return report_md
        return f"{lines[0].strip()}\n\n[구성대비]\n\n{chr(10).join(lines[1:]).strip()}"
    header = _rejection_basis_header(chain_info, matches)
    body = f"{header}\n\n[구성대비]\n\n{report_md}" if report_md else f"{header}\n\n[구성대비]"
    return body


def _build_related_inventions_tab(
    claim: ParsedClaim,
    prior_docs: List[ExtractedDocument],
    chain_info: Optional[dict],
    job_dir: Path,
) -> str:
    return build_rejected_inventions_section(
        claim,
        prior_docs,
        chain_info,
        str(job_dir),
    )


def _load_settings_with_dir():
    return load_settings()


def _job_dir(job_id: str) -> Path:
    return UPLOADS_DIR / job_id


def _case_dir(job_id: str) -> Path:
    return CASES_DIR / job_id


def _ensure_case_dirs(job_id: str) -> Path:
    case_dir = _case_dir(job_id)
    for name in ("pdfs", "parsed", "chunks", "vector_db", "reports"):
        (case_dir / name).mkdir(parents=True, exist_ok=True)
    return case_dir


def _save_chain_audit(job_id: str, chain_data: dict) -> None:
    """작업 폴더가 정리된 뒤에도 문헌 선정 이유를 사건별로 재현할 수 있게 보존한다."""
    case_dir = _ensure_case_dirs(job_id)
    _write_json(case_dir / "parsed" / "citation_chain.json", chain_data)
    job_dir = _job_dir(job_id)
    comparisons = {}
    for path in sorted(job_dir.glob("comparisons_*.json")):
        comparisons[path.stem] = _load_json(path, {})
    _write_json(case_dir / "parsed" / "comparison_matrix.json", comparisons)


def _file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _doc_cache_path(sha256: str) -> Path:
    return DOC_CACHE_DIR / f"{sha256}.json"


def _load_json(path: Path, default):
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def _write_json(path: Path, data) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _ordered_pdf_paths(job_dir: Path) -> list[Path]:
    """Return uploaded PDFs in the original upload order when possible.

    Falling back to filename order keeps backward compatibility for old jobs
    that do not have an upload manifest.
    """
    pdf_dir = job_dir / "pdfs"
    manifest = _load_json(job_dir / "upload_manifest.json", {})
    files = manifest.get("files") if isinstance(manifest, dict) else None
    ordered: list[Path] = []
    seen: set[Path] = set()

    if isinstance(files, list):
        for item in files:
            if not isinstance(item, dict):
                continue
            filename = Path(str(item.get("filename", "")).strip()).name
            if not filename:
                continue
            path = pdf_dir / filename
            if path.exists() and path not in seen:
                ordered.append(path)
                seen.add(path)

    for path in sorted(pdf_dir.glob("*.pdf")):
        if path not in seen:
            ordered.append(path)
    return ordered


def _report_timing_path(job_id: str, claim_number: int) -> Path:
    return _case_dir(job_id) / "reports" / f"claim{claim_number}_timing.json"


def _dependent_batch_status_path(job_id: str) -> Path:
    return _job_dir(job_id) / "dependent_batch_status.json"


def _update_dependent_batch_status(
    job_id: str,
    *,
    state: str,
    claim_numbers: list[int],
    stage: str,
    message: str,
    started_at: Optional[str] = None,
    error: str = "",
    reports_ready: int = 0,
    completed_at: Optional[str] = None,
) -> dict:
    path = _dependent_batch_status_path(job_id)
    previous = _load_json(path, {})
    now = time.strftime("%Y-%m-%d %H:%M:%S")
    payload = {
        "job_id": job_id,
        "state": state,
        "stage": stage,
        "message": message,
        "claim_numbers": claim_numbers,
        "reports_ready": reports_ready,
        "started_at": started_at or previous.get("started_at") or now,
        "updated_at": now,
        "completed_at": completed_at or previous.get("completed_at") or "",
        "error": error,
    }
    _write_json(path, payload)
    return payload


def _claim_analysis_signature(claim_data: Optional[dict]) -> tuple:
    """Return the claim fields that can change citation selection or inheritance."""
    if not isinstance(claim_data, dict):
        return ()
    elements = []
    for element in claim_data.get("elements") or []:
        if not isinstance(element, dict):
            continue
        elements.append((
            str(element.get("label") or ""),
            " ".join(str(element.get("text") or "").split()),
            str(element.get("importance") or "3"),
            bool(element.get("is_sub")),
            str(element.get("parent_label") or ""),
        ))
    return (
        int(claim_data.get("claim_number") or 0),
        str(claim_data.get("claim_type") or "independent"),
        claim_data.get("parent_claim"),
        " ".join(str(claim_data.get("text") or "").split()),
        " ".join(str(claim_data.get("preamble") or "").split()),
        " ".join(str(claim_data.get("closing") or "").split()),
        tuple(elements),
    )


def _invalidate_claim_derived_artifacts(
    job_id: str,
    claim_number: int,
    *,
    is_new_claim: bool = False,
    claim_type: str = "independent",
) -> None:
    """Invalidate only the changed claim family state that is no longer valid.

    Adding a new dependent claim must not erase an already issued independent-claim
    report, its context, citation chain, or document numbering. Existing descendants
    are invalidated when their ancestor text changes, while unrelated claim families
    remain intact.
    """
    job_dir = _job_dir(job_id)
    claim_key = str(claim_number)
    claims_data = _load_json(job_dir / "claims.json", [])
    affected_claim_numbers = {claim_number}
    changed = True
    while changed:
        changed = False
        for item in claims_data:
            try:
                number = int(item.get("claim_number"))
                parent = int(item.get("parent_claim")) if item.get("parent_claim") is not None else None
            except (TypeError, ValueError):
                continue
            if parent in affected_claim_numbers and number not in affected_claim_numbers:
                affected_claim_numbers.add(number)
                changed = True

    # A comparison file contains results for several claims. Preserve the other
    # claims and remove only the entry whose source text/elements changed.
    for path in job_dir.glob("comparisons_*.json"):
        cache = _load_json(path, None)
        if not isinstance(cache, dict) or claim_key not in cache:
            continue
        cache.pop(claim_key, None)
        _write_json(path, cache)

    # Category aliases are recomputed, but citation-chain numbering and unrelated
    # report context are stable case state and must survive a newly added claim.
    try:
        (job_dir / "same_pairs.json").unlink(missing_ok=True)
    except OSError as exc:
        logger.warning("Could not remove stale category alias file: %s", exc)

    context_path = job_dir / "context.json"
    context = _load_json(context_path, [])
    if isinstance(context, list):
        filtered_context = [
            item for item in context
            if item.get("claim_number") not in affected_claim_numbers
        ]
        if filtered_context != context:
            _write_json(context_path, filtered_context)

    # Editing an existing independent claim permits that family's references to be
    # selected again. New claims and dependent-claim edits preserve the root lock.
    chain_path = job_dir / "citation_chain.json"
    chain_data = _load_json(chain_path, {})
    if (
        not is_new_claim
        and claim_type == "independent"
        and isinstance(chain_data, dict)
        and isinstance(chain_data.get("selection_locks"), dict)
    ):
        locks = dict(chain_data.get("selection_locks") or {})
        if locks.pop(claim_key, None) is not None:
            chain_data["selection_locks"] = locks
            _write_json(chain_path, chain_data)

    report_paths = []
    for affected_number in affected_claim_numbers:
        report_paths.extend(REPORTS_DIR.glob(f"report_{job_id}_claim{affected_number}.*"))
    case_reports_dir = _case_dir(job_id) / "reports"
    if case_reports_dir.exists():
        for affected_number in affected_claim_numbers:
            report_paths.extend(case_reports_dir.glob(f"claim{affected_number}.*"))
    for path in report_paths:
        try:
            path.unlink(missing_ok=True)
        except OSError as exc:
            logger.warning("Could not remove stale report %s: %s", path, exc)

    for affected_number in affected_claim_numbers:
        _remove_reference_entries_for_claim(job_id, affected_number)


def _lock_claim_chain(job_id: str, claim_number: int) -> None:
    """Freeze an issued independent claim's references and document numbering."""
    chain_path = _job_dir(job_id) / "citation_chain.json"
    chain_data = _load_json(chain_path, {})
    if not isinstance(chain_data, dict):
        return
    chain = (chain_data.get("chains") or {}).get(str(claim_number)) or {}
    total = [int(value) for value in chain.get("total", [])]
    if not total:
        return
    locks = dict(chain_data.get("selection_locks") or {})
    locks[str(claim_number)] = {
        "total": total,
        "doc_name_mapping": dict(chain_data.get("doc_name_mapping") or {}),
        "locked_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "reason": "independent_claim_report_issued",
    }
    chain_data["selection_locks"] = locks
    _write_json(chain_path, chain_data)
    _save_chain_audit(job_id, chain_data)


def _load_claims(job_id: str) -> List[ParsedClaim]:
    data = _load_json(_job_dir(job_id) / "claims.json", [])
    return [ParsedClaim(**item) for item in data]


def _load_prior_docs(job_id: str) -> List[ExtractedDocument]:
    data = _load_json(_job_dir(job_id) / "prior_docs.json", [])
    return [ExtractedDocument(**item) for item in data]


def _load_context(job_id: str) -> list:
    return _load_json(_job_dir(job_id) / "context.json", [])


def _save_context_entry(job_id: str, claim_number: int, claim_text: str, report_md: str) -> None:
    context = [c for c in _load_context(job_id) if c.get("claim_number") != claim_number]
    summary = _phase1_only_report(report_md)[-4000:]
    context.append({
        "claim_number": claim_number,
        "claim_text_preview": claim_text[:200],
        "report_summary": summary,
    })
    context.sort(key=lambda x: x.get("claim_number", 0))
    _write_json(_job_dir(job_id) / "context.json", context)


def _parent_chain_nums(claim: ParsedClaim, claims_by_num: dict[int, ParsedClaim]) -> list[int]:
    """Return direct-to-root parent claim numbers for a dependent claim."""
    chain: list[int] = []
    cur = claim
    visited: set[int] = set()
    while cur and cur.claim_type == "dependent" and cur.parent_claim:
        parent_num = cur.parent_claim
        if parent_num in visited:
            break
        visited.add(parent_num)
        chain.append(parent_num)
        cur = claims_by_num.get(parent_num)
    return chain


def _context_for_claims(job_id: str, claim_numbers: set[int]) -> list:
    if not claim_numbers:
        return []
    selected = [
        c for c in _load_context(job_id)
        if c.get("claim_number") in claim_numbers
    ]
    selected.sort(key=lambda x: x.get("claim_number", 0))
    return selected


def _parent_independent_num(claim: ParsedClaim, claims_by_num: dict[int, ParsedClaim]) -> Optional[int]:
    if claim.claim_type != "dependent":
        return None
    cur = claim
    visited: set[int] = set()
    while cur and cur.claim_type == "dependent" and cur.parent_claim:
        if cur.claim_number in visited:
            return None
        visited.add(cur.claim_number)
        cur = claims_by_num.get(cur.parent_claim)
    return cur.claim_number if cur and cur.claim_type == "independent" else None


def _save_case_artifacts(job_id: str, docs: List[ExtractedDocument], manifest: list[dict]) -> None:
    case_dir = _ensure_case_dirs(job_id)
    _write_json(case_dir / "case_metadata.json", {
        "case_id": job_id,
        "prior_count": len(docs),
        "created_or_updated": time.strftime("%Y-%m-%d %H:%M:%S"),
        "scope": "업로드된 파일의 대응관계만 해당 케이스 한정으로 사용",
        "manifest": manifest,
    })
    _write_json(case_dir / "parsed" / "prior_docs.json", [d.model_dump() for d in docs])
    _write_json(case_dir / "parsed" / "paragraph_records.json", [
        rec.model_dump() for doc in docs for rec in doc.paragraph_records
    ])
    _write_json(case_dir / "chunks" / "paragraph_chunks.json", [
        chunk.model_dump() for doc in docs for chunk in doc.paragraph_chunks
    ])
    _write_json(case_dir / "chunks" / "group_chunks.json", [
        chunk.model_dump() for doc in docs for chunk in doc.group_chunks
    ])
    save_case_artifacts_sqlite(case_dir, docs, manifest)


def _save_report(job_id: str, claim_number: int, md: str) -> None:
    md = sanitize_report_status_icons(md)
    path = REPORTS_DIR / f"report_{job_id}_claim{claim_number}.md"
    path.write_text(md, encoding="utf-8")
    case_report = _ensure_case_dirs(job_id) / "reports" / f"claim{claim_number}.md"
    case_report.write_text(md, encoding="utf-8")


def _save_reference_db(
    job_id: str,
    claim: ParsedClaim,
    matches,
    prior_docs: List[ExtractedDocument],
    chain_info: Optional[dict],
    report_md: str,
) -> None:
    used = (chain_info or {}).get("total") or sorted({m.cited_invention_index for m in matches})
    chain_roles = (chain_info or {}).get("reference_roles") or {}
    persisted_role_names = {
        "primary": "primary_reference",
        "substantive_secondary": "secondary_reference",
        "conventional_support": "conventional_support",
    }
    role_by_idx = {
        idx: persisted_role_names.get(
            chain_roles.get(str(idx), ""),
            "primary_reference" if order == 0 else "secondary_reference",
        )
        for order, idx in enumerate(used)
    }
    is_novelty = _is_novelty_basis(chain_info, matches, list(used))
    entries = []
    for doc_idx in used:
        if doc_idx < 0 or doc_idx >= len(prior_docs):
            continue
        doc = prior_docs[doc_idx]
        doc_matches = [m for m in matches if m.cited_invention_index == doc_idx and m.quote]
        entries.append({
            "publication_no": doc.publication_no or doc.filename,
            "title": doc.title or doc.filename,
            "used_in_case": job_id,
            "claim_number": claim.claim_number,
            "role": role_by_idx.get(doc_idx, "reference"),
            "rejection_type": "novelty" if is_novelty else "inventive_step",
            "key_paragraphs": [m.chunk_id.strip("[]") for m in doc_matches],
            "matched_features": [m.label for m in doc_matches],
            "report_excerpt": [
                {"paragraph_no": m.chunk_id.strip("[]"), "quote": m.quote}
                for m in doc_matches
            ],
        })
    if not entries:
        _remove_reference_entries_for_claim(job_id, claim.claim_number)
        return
    case_path = _ensure_case_dirs(job_id) / "reference_db.json"
    case_entries = _load_json(case_path, [])
    case_entries = [
        item for item in case_entries
        if not (item.get("used_in_case") == job_id and item.get("claim_number") == claim.claim_number)
    ]
    case_entries.extend(entries)
    _write_json(case_path, case_entries)
    save_reference_entries_sqlite(_ensure_case_dirs(job_id), entries)
    cumulative_path = CASES_DIR / "reference_db.json"
    cumulative = _load_json(cumulative_path, [])
    cumulative = [
        item for item in cumulative
        if not (item.get("used_in_case") == job_id and item.get("claim_number") == claim.claim_number)
    ]
    cumulative.extend(entries)
    _write_json(cumulative_path, cumulative)
    save_reference_entries_sqlite(CASES_DIR, entries)


def _remove_reference_entries_for_claim(job_id: str, claim_number: int) -> None:
    for json_path in (
        _case_dir(job_id) / "reference_db.json",
        CASES_DIR / "reference_db.json",
    ):
        if not json_path.exists():
            continue
        entries = _load_json(json_path, [])
        filtered = [
            item for item in entries
            if not (
                item.get("used_in_case") == job_id
                and item.get("claim_number") == claim_number
            )
        ]
        _write_json(json_path, filtered)

    for db_path in (
        _case_dir(job_id) / "reference.sqlite",
        CASES_DIR / "reference.sqlite",
    ):
        if not db_path.exists():
            continue
        try:
            with sqlite3.connect(str(db_path)) as conn:
                conn.execute(
                    "DELETE FROM reference_entries WHERE used_in_case = ? AND claim_number = ?",
                    (job_id, claim_number),
                )
        except sqlite3.Error:
            logger.warning(
                "Failed to prune reference_entries for %s claim %s",
                job_id,
                claim_number,
                exc_info=True,
            )


def _remove_reference_entries_for_job(job_id: str) -> None:
    cumulative_path = CASES_DIR / "reference_db.json"
    if cumulative_path.exists():
        cumulative = _load_json(cumulative_path, [])
        cumulative = [item for item in cumulative if item.get("used_in_case") != job_id]
        _write_json(cumulative_path, cumulative)

    db_path = CASES_DIR / "reference.sqlite"
    if db_path.exists():
        try:
            with sqlite3.connect(str(db_path)) as conn:
                conn.execute("DELETE FROM reference_entries WHERE used_in_case = ?", (job_id,))
        except sqlite3.Error:
            logger.warning("Failed to prune reference_entries for %s", job_id, exc_info=True)


def _delete_doc_cache_for_job(job_id: str) -> None:
    shas: set[str] = set()
    for path in (
        _job_dir(job_id) / "prior_manifest.json",
        _case_dir(job_id) / "case_metadata.json",
    ):
        data = _load_json(path, {})
        items = data.get("manifest") if isinstance(data, dict) else data
        if not isinstance(items, list):
            continue
        for item in items:
            sha = item.get("sha256") if isinstance(item, dict) else None
            if sha:
                shas.add(str(sha))

    for sha in shas:
        (_doc_cache_path(sha)).unlink(missing_ok=True)


def _rmtree_with_retry(path: Path, attempts: int = 5) -> None:
    last_error: Exception | None = None
    for _ in range(attempts):
        try:
            shutil.rmtree(path)
            return
        except PermissionError as exc:
            last_error = exc
            gc.collect()
            time.sleep(0.2)
    if last_error:
        raise last_error


@router.post("/upload")
async def upload(prior_files: List[UploadFile], base_job_id: Optional[str] = Form(default=None)):
    if len(prior_files) > 7:
        raise HTTPException(status_code=400, detail="인용발명 PDF는 최대 7개까지 업로드할 수 있습니다.")

    job_id = f"CASE-{time.strftime('%Y%m%d')}-{uuid.uuid4().hex[:8]}"
    job_dir = _job_dir(job_id)
    pdf_dir = job_dir / "pdfs"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    case_dir = _ensure_case_dirs(job_id)

    saved = []
    for file in prior_files:
        filename = Path(file.filename or "prior.pdf").name
        dest = pdf_dir / filename
        async with aiofiles.open(dest, "wb") as out:
            while chunk := await file.read(1024 * 1024):
                await out.write(chunk)
        shutil.copy2(dest, case_dir / "pdfs" / filename)
        saved.append({"filename": filename, "path": str(dest)})

    _write_json(job_dir / "upload_manifest.json", {
        "job_id": job_id,
        "base_job_id": base_job_id,
        "files": saved,
    })
    return {"job_id": job_id, "files": saved}


@router.get("/prepare/{job_id}")
async def prepare(job_id: str):
    async def stream() -> AsyncGenerator[dict, None]:
        job_dir = _job_dir(job_id)
        if not job_dir.exists():
            yield _ev("error", "작업을 찾을 수 없습니다.")
            return
        pdfs = _ordered_pdf_paths(job_dir)
        if not pdfs:
            yield _ev("error", "업로드된 PDF가 없습니다.")
            return

        docs: list[ExtractedDocument] = []
        manifest: list[dict] = []
        for idx, pdf_path in enumerate(pdfs):
            try:
                yield _ev("extract_prior", f"{pdf_path.name} 텍스트 추출 중..")
                sha = _file_sha256(pdf_path)
                cache_path = _doc_cache_path(sha)
                if cache_path.exists():
                    doc = ExtractedDocument(**_load_json(cache_path, {}))
                    resolved_doc_id = f"D{idx + 1}"
                    doc = doc.model_copy(update={
                        "doc_index": idx,
                        "doc_id": resolved_doc_id,
                        "pdf_path": str(pdf_path.resolve()),
                        "filename": pdf_path.name,
                        "paragraph_records": [
                            rec.model_copy(update={"doc_id": resolved_doc_id})
                            for rec in (doc.paragraph_records or [])
                        ],
                        "paragraph_chunks": [
                            chunk.model_copy(update={"doc_id": resolved_doc_id})
                            for chunk in (doc.paragraph_chunks or [])
                        ],
                        "group_chunks": [
                            chunk.model_copy(update={"doc_id": resolved_doc_id})
                            for chunk in (doc.group_chunks or [])
                        ],
                    })
                    yield _ev("extract_prior", f"{pdf_path.name} 로드 완료")
                else:
                    doc = pdf_extractor.extract(str(pdf_path), idx)
                    _write_json(cache_path, doc.model_dump())
                docs.append(doc)
                manifest.append({
                    "doc_id": doc.doc_id or f"D{idx + 1}",
                    "filename": pdf_path.name,
                    "sha256": sha,
                    "publication_no": doc.publication_no,
                    "title": doc.title,
                    "paragraph_count": len(doc.paragraph_records or doc.paragraphs),
                    "paragraph_chunk_count": len(doc.paragraph_chunks),
                    "group_chunk_count": len(doc.group_chunks),
                })
                yield _ev("extract_prior_done", f"{pdf_path.name}: 추출 {manifest[-1]['paragraph_count']}개")
            except Exception as e:
                yield _ev("error", f"{pdf_path.name} 추출 실패: {e}")
                return

        _write_json(job_dir / "prior_docs.json", [d.model_dump() for d in docs])
        _write_json(job_dir / "prior_manifest.json", manifest)
        _save_case_artifacts(job_id, docs, manifest)
        yield _ev("prepare_done", {
            "job_id": job_id,
            "prior_count": len(docs),
            "manifest": manifest,
        })

    return EventSourceResponse(stream(), headers={"Content-Type": "text/event-stream; charset=utf-8"})


@router.post("/manual_claim/{job_id}")
async def manual_claim(job_id: str, req: ManualClaimRequest):
    job_dir = _job_dir(job_id)
    if not job_dir.exists():
        raise HTTPException(status_code=404, detail="작업을 찾을 수 없습니다.")
    if req.parent_claim == req.claim_number:
        raise HTTPException(
            status_code=422,
            detail="종속항은 자기 자신을 부모 청구항으로 참조할 수 없습니다.",
        )
    claim = await parse_manual_claim_locally(
        req.claim_text,
        req.claim_number,
        req.claim_type,
        req.parent_claim,
    )
    stored_claims = _load_json(job_dir / "claims.json", [])
    claim_data = claim.model_dump()
    previous = next((c for c in stored_claims if c.get("claim_number") == claim.claim_number), None)
    claims = [c for c in stored_claims if c.get("claim_number") != claim.claim_number]
    claims.append(claim_data)
    claims.sort(key=lambda c: c.get("claim_number", 0))
    _write_json(job_dir / "claims.json", claims)
    _write_json(_ensure_case_dirs(job_id) / "parsed" / "claims.json", claims)
    if _claim_analysis_signature(previous) != _claim_analysis_signature(claim_data):
        _invalidate_claim_derived_artifacts(
            job_id,
            claim.claim_number,
            is_new_claim=previous is None,
            claim_type=claim.claim_type,
        )
    return claim.model_dump()


@router.post("/detect_category/{job_id}")
async def detect_category(job_id: str):
    settings = _load_settings_with_dir()
    claims = _load_claims(job_id)
    same_pairs = await detect_category_same_claims(claims, settings)
    _write_json(_job_dir(job_id) / "same_pairs.json", same_pairs)
    return {"same_pairs": same_pairs}


@router.get("/report/{job_id}/{claim_number}")
async def report(job_id: str, claim_number: int, use_context: bool = True, force: bool = False):
    async def stream() -> AsyncGenerator[dict, None]:
        try:
            total_start = time.perf_counter()
            timing_data: dict[str, float] = {}

            def _timing_message(label: str, started_at: float) -> str:
                return f"[timing] {label}: {_elapsed(started_at)}"

            async def _yield_timing(label: str, started_at: float) -> AsyncGenerator[dict, None]:
                timing_data[label] = round(time.perf_counter() - started_at, 3)
                message = _timing_message(label, started_at)
                logger.info("Report timing [%s/%s] %s", job_id, claim_number, message)
                yield _ev("log", message)

            job_dir = _job_dir(job_id)
            if not job_dir.exists():
                yield _ev("error", "작업을 찾을 수 없습니다.")
                return
            settings = _load_settings_with_dir()
            compare_mode = getattr(settings, "comparison_mode", "per_doc")

            cached = REPORTS_DIR / f"report_{job_id}_claim{claim_number}.md"
            claims = _load_claims(job_id)
            prior_docs = _load_prior_docs(job_id)
            same_pairs = _load_json(job_dir / "same_pairs.json", {})
            claim = next((c for c in claims if c.claim_number == claim_number), None)
            if not claim:
                yield _ev("error", f"청구항 {claim_number}를 찾을 수 없습니다.")
                return
            if not prior_docs:
                yield _ev("error", "인용발명이 준비되지 않았습니다.")
                return

            cache_reset = reset_incompatible_comparison_caches(
                str(job_dir), len(prior_docs), settings
            )
            matches, cached_all = get_matches_from_cache(
                claim,
                prior_docs,
                str(job_dir),
                comparison_mode=compare_mode,
            )
            cached_chain = _load_json(job_dir / "citation_chain.json", {})
            policy_cache_current = (
                isinstance(cached_chain, dict)
                and cached_chain.get("policy_version") == CITATION_CHAIN_POLICY_VERSION
            )
            if (
                cached.exists()
                and not force
                and cached_all
                and policy_cache_current
            ):
                cached_report = sanitize_report_status_icons(_phase1_only_report(cached.read_text(encoding="utf-8")))
                cached_chain_info = get_claim_chain_info(cached_chain, claim_number)
                cached_scope_violations = find_unselected_reference_mentions(
                    cached_report, cached_chain_info
                )
                if not cached_scope_violations:
                    _save_context_entry(job_id, claim_number, claim.text, cached_report)
                    if claim.claim_type == "independent":
                        _lock_claim_chain(job_id, claim_number)
                    async for event in _yield_timing("cached report return", total_start):
                        yield event
                    yield _ev("done", {
                        "report_md": cached_report,
                        "claim_number": claim_number,
                        "used_inventions": _used_inventions_for(
                            cached_chain_info, prior_docs, matches=matches
                        ),
                    })
                    return
                yield _ev(
                    "log",
                    "[cache] 최종 인용 체인과 다른 기존 보고서를 무효화하고 재작성합니다: "
                    + ", ".join(cached_scope_violations),
                )

            yield _ev("start", f"청구항 {claim_number} 보고서 작성을 시작합니다.")
            if cache_reset:
                yield _ev("log", "[cache] incompatible comparison cache reset")
            if not cached_all:
                compare_start = time.perf_counter()
                cached_doc_idxs = get_cached_doc_indices(
                    str(job_dir),
                    claim_number,
                    len(prior_docs),
                    comparison_mode=compare_mode,
                )
                missing_doc_idxs = [i for i in range(len(prior_docs)) if i not in cached_doc_idxs]
                try:
                    if cached_doc_idxs and missing_doc_idxs:
                        yield _ev("analyze", f"missing comparison docs: {len(missing_doc_idxs)}")
                        await analyze_claim_elements_for_docs(
                            claim.elements, prior_docs, missing_doc_idxs, settings,
                            job_dir=str(job_dir), claim_number=claim_number,
                        )
                    else:
                        if compare_mode in {"mixed", "hybrid"}:
                            yield _ev(
                                "analyze",
                                f"comparing {len(prior_docs)} prior docs in integrated mode",
                            )
                        else:
                            yield _ev(
                                "analyze",
                                f"comparing {len(prior_docs)} prior docs in per-doc mode",
                            )
                        compare_fn = analyze_claim_elements if compare_mode == "per_doc" else analyze_claim_elements_hybrid
                        await compare_fn(
                            claim.elements, prior_docs, settings,
                            job_dir=str(job_dir), claim_number=claim_number,
                        )
                except CompareFailed as e:
                    yield _ev("error", f"구성요소 비교 분석 실패: {e}")
                    return
                matches, _ = get_matches_from_cache(
                    claim,
                    prior_docs,
                    str(job_dir),
                    comparison_mode=compare_mode,
                )
                async for event in _yield_timing("comparison", compare_start):
                    yield event
            else:
                yield _ev("log", "[cache] using cached comparisons")


            chain_start = time.perf_counter()
            chain_data = build_citation_chain_from_comparisons(str(job_dir), claims, prior_docs)
            _save_chain_audit(job_id, chain_data)
            async for event in _yield_timing("citation chain", chain_start):
                yield event
            chain_info = get_claim_chain_info(chain_data, claim_number) if chain_data else None
            if chain_info and chain_info.get("total"):
                matches, _ = get_matches_from_cache(
                    claim,
                    prior_docs,
                    str(job_dir),
                    allowed_docs=chain_info["total"],
                    comparison_mode=compare_mode,
                )

            secondary_matches = None
            total_refs = (chain_info or {}).get("total", [])
            report_matches = matches
            if claim.claim_type == "independent" and total_refs:
                report_matches, _ = get_matches_from_cache(
                    claim,
                    prior_docs,
                    str(job_dir),
                    allowed_docs=[total_refs[0]],
                    comparison_mode=compare_mode,
                )
            if len(total_refs) > 1:
                secondary_matches = []
                for sec_idx in total_refs[1:]:
                    sec_matches, _ = get_matches_from_cache(
                        claim,
                        prior_docs,
                        str(job_dir),
                        allowed_docs=[sec_idx],
                        comparison_mode=compare_mode,
                    )
                    secondary_matches.extend(sec_matches)

            yield _ev("log", "[verify] checking quote text against local DB")
            verify_start = time.perf_counter()
            verifications = verify_quotes(matches, prior_docs)
            async for event in _yield_timing("quote verification", verify_start):
                yield event
            for item in verifications:
                yield _ev("log", item["message"])

            prev_context = []
            if use_context and claim.claim_type == "dependent":
                claims_by_num = {c.claim_number: c for c in claims}
                parent_num = _parent_independent_num(claim, claims_by_num)
                if parent_num is not None:
                    prev_context = [c for c in _load_context(job_id) if c.get("claim_number") == parent_num]

            used_inventions = _used_inventions_for(
                chain_info, prior_docs, matches=report_matches
            )
            yield _ev("generate", "Phase 1 analysis in progress")
            phase1_start = time.perf_counter()
            phase1_chunks: list[str] = []
            if claim.claim_type == "independent":
                async for chunk in generate_independent_phase1_streaming(
                    claim, report_matches, prior_docs, chain_info, settings,
                    prev_context=prev_context,
                    secondary_matches=secondary_matches,
                ):
                    clean_chunk = sanitize_report_status_icons(chunk)
                    phase1_chunks.append(clean_chunk)
                    yield _ev("stream_chunk", clean_chunk)
                phase1_md = polish_phase1_summary_text(
                    sanitize_report_status_icons(_dedupe_phase1_sections(_strip_agent_tool_calls("".join(phase1_chunks))))
                )
                phase1_md = enforce_phase1_judgment_headers(phase1_md, report_matches)
            else:
                async def _emit_progress(message: str) -> None:
                    yield_event = _ev("log", message)
                    nonlocal_yield_events.append(yield_event)

                nonlocal_yield_events: list[dict] = []
                raw = await _await_with_log_heartbeat(
                    generate_dependent_report(
                        claim, matches, prior_docs, chain_info, settings,
                        prev_context=prev_context,
                        secondary_matches=secondary_matches,
                    ),
                    _emit_progress,
                    label=f"청구항 {claim_number} Phase 1 작성",
                )
                for event in nonlocal_yield_events:
                    yield event
                raw = sanitize_report_status_icons(_strip_agent_tool_calls(raw))
                split = _find_legacy_final_boundary(raw)
                phase1_body = raw[:split].strip() if split >= 0 else raw
                phase1_md = polish_phase1_summary_text(
                    sanitize_report_status_icons(_dedupe_phase1_sections(phase1_body))
                )

            scope_violations = find_unselected_reference_mentions(phase1_md, chain_info)
            if scope_violations and claim.claim_type == "independent":
                yield _ev(
                    "log",
                    "[verify] 최종 인용 체인 밖의 문헌이 본문에 포함되어 보고서를 다시 작성합니다: "
                    + ", ".join(scope_violations),
                )
                strict_chain_info = dict(chain_info or {})
                strict_chain_info["strict_reference_scope_retry"] = True
                retry_chunks: list[str] = []
                async for chunk in generate_independent_phase1_streaming(
                    claim, report_matches, prior_docs, strict_chain_info, settings,
                    prev_context=prev_context,
                    secondary_matches=secondary_matches,
                ):
                    retry_chunks.append(sanitize_report_status_icons(chunk))
                phase1_md = polish_phase1_summary_text(
                    sanitize_report_status_icons(
                        _dedupe_phase1_sections(_strip_agent_tool_calls("".join(retry_chunks)))
                    )
                )
                phase1_md = enforce_phase1_judgment_headers(phase1_md, report_matches)
                scope_violations = find_unselected_reference_mentions(phase1_md, chain_info)
            if scope_violations:
                raise CompareFailed(
                    "보고서가 최종 인용 체인에 포함되지 않은 문헌을 사용했습니다: "
                    + ", ".join(scope_violations)
                )
            phase1_md = f"### claim {claim_number}\n\n{phase1_md}"
            async for event in _yield_timing("phase1", phase1_start):
                yield event
            yield _ev("phase1_result", {
                "phase1_md": phase1_md,
                "claim_number": claim_number,
                "used_inventions": used_inventions,
            })

            report_md = _prepend_rejection_basis_header(
                sanitize_report_status_icons(phase1_md),
                chain_info,
                matches,
            )
            finalize_start = time.perf_counter()
            same_claims_for_this = [int(k) for k, v in same_pairs.items() if v == claim_number]
            if same_claims_for_this:
                report_md = sanitize_report_status_icons(generate_category_same_report(claim_number, same_claims_for_this, report_md))
            related_inventions_md = _build_related_inventions_tab(
                claim,
                prior_docs,
                chain_info,
                job_dir,
            )

            _save_report(job_id, claim_number, report_md)
            _save_reference_db(job_id, claim, matches, prior_docs, chain_info, report_md)
            _save_context_entry(job_id, claim_number, claim.text, report_md)
            if claim.claim_type == "independent":
                _lock_claim_chain(job_id, claim_number)
            async for event in _yield_timing("finalize", finalize_start):
                yield event
            async for event in _yield_timing("total", total_start):
                yield event
            _write_json(
                _report_timing_path(job_id, claim_number),
                {
                    "job_id": job_id,
                    "claim_number": claim_number,
                    "timings": timing_data,
                },
            )
            yield _ev("done", {
                "report_md": report_md,
                "related_inventions_md": related_inventions_md,
                "claim_number": claim_number,
                "used_inventions": used_inventions,
                "timings": timing_data,
            })
        except Exception as e:
            tb = traceback.format_exc()
            logger.error(f"Report error [{job_id}/{claim_number}]: {tb}")
            yield _ev("error", f"error: {e}\n{tb[:500]}")

    return EventSourceResponse(stream(), headers={"Content-Type": "text/event-stream; charset=utf-8"})


def _used_inventions_for(
    chain_info,
    prior_docs: List[ExtractedDocument],
    *,
    matches=None,
) -> list:
    if chain_info:
        mapping = chain_info.get("doc_name_mapping", {})
        records = [
            {
                "name": mapping.get(str(idx), f"인용발명 {idx + 1}"),
                "filename": prior_docs[idx].filename if idx < len(prior_docs) else "",
                "role": (chain_info.get("reference_roles") or {}).get(str(idx), ""),
            }
            for idx in chain_info.get("total", [])
        ]
        if records:
            records[0]["basis_label"] = _rejection_basis_header(
                chain_info, list(matches or [])
            ).strip("[]")
        return records
    return [{"name": "인용발명 1", "filename": prior_docs[0].filename}] if prior_docs else []


def _assemble_dependent_report(
    raw: str,
    claim: ParsedClaim,
    chain_info,
    settings,
    matches=None,
    secondary_matches=None,
) -> str:
    body = _strip_agent_tool_calls(raw)
    split = _find_legacy_final_boundary(body)
    phase1 = body[:split].strip() if split >= 0 else body.strip()
    phase1 = _dedupe_phase1_sections(phase1)
    phase1 = enforce_phase1_judgment_headers(phase1, list(matches or []))
    return f"### 청구항 {claim.claim_number}\n\n{phase1}"


def _has_substantive_dependent_report(md: str, claim_number: int) -> bool:
    body = sanitize_report_status_icons(_strip_agent_tool_calls(md or "")).strip()
    body = re.sub(r"(?m)^\s*\[[^\]\n]*(?:신규성|진보성)[^\]\n]*\]\s*$", "", body)
    body = re.sub(r"(?m)^\s*\[구성대비\]\s*$", "", body)
    body = re.sub(rf"(?mi)^\s*###\s*(?:청구항|claim)\s*{claim_number}\s*$", "", body)
    body = re.sub(r"(?m)^\s*[-*_]{3,}\s*$", "", body).strip()
    return len(body) >= 20


@router.post("/report_batch_dependent/{job_id}")
async def report_batch_dependent(job_id: str, req: BatchDependentRequest):
    job_dir = _job_dir(job_id)
    if not job_dir.exists():
        raise HTTPException(status_code=404, detail="작업을 찾을 수 없습니다.")

    claim_numbers = sorted({int(n) for n in req.claim_numbers})
    started_at = time.strftime("%Y-%m-%d %H:%M:%S")
    _update_dependent_batch_status(
        job_id,
        state="running",
        claim_numbers=claim_numbers,
        stage="starting",
        message="종속항 배치 보고서를 생성을 시작합니다.",
        started_at=started_at,
    )

    try:
        settings = _load_settings_with_dir()
        claims = _load_claims(job_id)
        prior_docs = _load_prior_docs(job_id)
        same_pairs = _load_json(job_dir / "same_pairs.json", {})
        claims_by_num = {c.claim_number: c for c in claims}
        targets = [
            claims_by_num[n] for n in req.claim_numbers
            if n in claims_by_num and claims_by_num[n].claim_type == "dependent"
        ]
        if not targets:
            _update_dependent_batch_status(
                job_id,
                state="completed",
                claim_numbers=claim_numbers,
                stage="completed",
                message="대상 종속항이 없어 빈 결과를 반환합니다.",
                started_at=started_at,
                completed_at=time.strftime("%Y-%m-%d %H:%M:%S"),
            )
            return {"reports": {}}

        _update_dependent_batch_status(
            job_id,
            state="running",
            claim_numbers=claim_numbers,
            stage="loaded_targets",
            message=f"종속항 {len(targets)}개를 배치 처리 대상으로 불러왔습니다.",
            started_at=started_at,
        )

        compare_mode = getattr(settings, "comparison_mode", "per_doc")

        cached_chain = _load_json(job_dir / "citation_chain.json", {})
        policy_cache_changed = (
            not isinstance(cached_chain, dict)
            or cached_chain.get("policy_version") != CITATION_CHAIN_POLICY_VERSION
        )
        reset_incompatible_comparison_caches(str(job_dir), len(prior_docs), settings)
        candidate_docs_by_claim: dict[int, list[int]] = {}

        def candidate_docs_for(claim: ParsedClaim) -> list[int]:
            if claim.claim_number not in candidate_docs_by_claim:
                candidate_docs_by_claim[claim.claim_number] = select_candidate_doc_indices_for_elements(
                    claim.elements,
                    prior_docs,
                    settings,
                )
            return candidate_docs_by_claim[claim.claim_number]

        async def ensure_comparison_cache(
            claim: ParsedClaim,
            *,
            force_refresh: bool = False,
        ) -> None:
            if str(claim.claim_number) in same_pairs:
                return
            cached_doc_idxs = get_cached_doc_indices(
                str(job_dir),
                claim.claim_number,
                len(prior_docs),
                comparison_mode=compare_mode,
            )
            target_doc_idxs = candidate_docs_for(claim)
            missing_doc_idxs = [i for i in target_doc_idxs if i not in cached_doc_idxs]
            if force_refresh:
                missing_doc_idxs = target_doc_idxs[:]
            if not missing_doc_idxs and not force_refresh:
                return
            if compare_mode in {"mixed", "hybrid"} and len(target_doc_idxs) > 1:
                selected_docs = [prior_docs[i] for i in target_doc_idxs]
                await analyze_claim_elements_hybrid(
                    claim.elements, selected_docs, settings,
                    job_dir=str(job_dir), claim_number=claim.claim_number,
                    doc_index_map=target_doc_idxs,
                )
                return
            if missing_doc_idxs:
                await analyze_claim_elements_for_docs(
                    claim.elements, prior_docs, missing_doc_idxs, settings,
                    job_dir=str(job_dir), claim_number=claim.claim_number,
                )
                return
            compare_fn = analyze_claim_elements if compare_mode == "per_doc" else analyze_claim_elements_hybrid
            await compare_fn(
                claim.elements, prior_docs, settings,
                job_dir=str(job_dir), claim_number=claim.claim_number,
            )

        ancestor_numbers: set[int] = set()
        for claim in targets:
            ancestor_numbers.update(_parent_chain_nums(claim, claims_by_num))
        ancestor_claims = [
            claims_by_num[number]
            for number in sorted(ancestor_numbers)
            if number in claims_by_num
        ]
        uncached_ancestors = []
        for claim in ancestor_claims:
            if str(claim.claim_number) in same_pairs:
                continue
            target_doc_idxs = candidate_docs_for(claim)
            cached_doc_idxs = get_cached_doc_indices(
                str(job_dir),
                claim.claim_number,
                len(prior_docs),
                comparison_mode=compare_mode,
            )
            if target_doc_idxs and not all(i in cached_doc_idxs for i in target_doc_idxs):
                uncached_ancestors.append(claim)

        if uncached_ancestors:
            _update_dependent_batch_status(
                job_id,
                state="running",
                claim_numbers=claim_numbers,
                stage="restoring_parent_comparison_cache",
                message=(
                    "부모항 인용 체인을 복구하기 위해 "
                    f"상위 청구항 {len(uncached_ancestors)}개의 비교 캐시를 생성하고 있습니다."
                ),
                started_at=started_at,
            )
            for ancestor in uncached_ancestors:
                await ensure_comparison_cache(ancestor, force_refresh=False)

        uncached_targets = []
        for claim in targets:
            if str(claim.claim_number) in same_pairs:
                continue
            _, dep_cached = get_matches_from_cache(
                claim,
                prior_docs,
                str(job_dir),
                comparison_mode=compare_mode,
            )
            target_doc_idxs = candidate_docs_for(claim)
            cached_doc_idxs = get_cached_doc_indices(
                str(job_dir),
                claim.claim_number,
                len(prior_docs),
                comparison_mode=compare_mode,
            )
            target_cached = target_doc_idxs and all(i in cached_doc_idxs for i in target_doc_idxs)
            if (dep_cached or target_cached) and not req.force:
                continue
            uncached_targets.append(claim)

        if uncached_targets:
            _update_dependent_batch_status(
                job_id,
                state="running",
                claim_numbers=claim_numbers,
                stage="building_comparison_cache",
                message=f"종속항 {len(uncached_targets)}개에 대해 비교 캐시를 생성하고 있습니다.",
                started_at=started_at,
            )
            max_parallel = 2 if compare_mode in {"mixed", "hybrid"} else 1
            semaphore = asyncio.Semaphore(max_parallel)

            async def run_limited(claim: ParsedClaim) -> None:
                async with semaphore:
                    await ensure_comparison_cache(claim, force_refresh=req.force)

            await asyncio.gather(*(run_limited(claim) for claim in uncached_targets))
        else:
            _update_dependent_batch_status(
                job_id,
                state="running",
                claim_numbers=claim_numbers,
                stage="comparison_cache_ready",
                message="비교 캐시가 이미 준비되어 있습니다.",
                started_at=started_at,
            )

        recomputed_ancestor_numbers = {
            claim.claim_number for claim in uncached_ancestors
        }
        recomputed_claim_numbers = {claim.claim_number for claim in uncached_targets}
        for claim in targets:
            if recomputed_ancestor_numbers.intersection(
                _parent_chain_nums(claim, claims_by_num)
            ):
                recomputed_claim_numbers.add(claim.claim_number)

        _update_dependent_batch_status(
            job_id,
            state="running",
            claim_numbers=claim_numbers,
            stage="building_citation_chain",
            message="인용발명 체인을 다시 계산하고 있습니다.",
            started_at=started_at,
        )
        chain_data = build_citation_chain_from_comparisons(str(job_dir), claims, prior_docs)
        _save_chain_audit(job_id, chain_data)
        parent_nums: set[int] = set()
        for claim in targets:
            parent_nums.update(_parent_chain_nums(claim, claims_by_num))
        prev_context = _context_for_claims(job_id, parent_nums)
        batch_items = []
        results: dict[str, dict] = {}

        for claim in targets:
            cn = claim.claim_number
            cached = REPORTS_DIR / f"report_{job_id}_claim{cn}.md"
            if (
                cached.exists()
                and not req.force
                and cn not in recomputed_claim_numbers
                and not policy_cache_changed
            ):
                cached_report = sanitize_report_status_icons(cached.read_text(encoding="utf-8"))
                if not _has_substantive_dependent_report(cached_report, cn):
                    cached.unlink(missing_ok=True)
                    (_ensure_case_dirs(job_id) / "reports" / f"claim{cn}.md").unlink(missing_ok=True)
                else:
                    cached_chain_info = get_claim_chain_info(chain_data, cn)
                    cached_matches, _ = get_matches_from_cache(
                        claim,
                        prior_docs,
                        str(job_dir),
                        allowed_docs=(cached_chain_info or {}).get("total"),
                        comparison_mode=compare_mode,
                    )
                    _save_context_entry(job_id, claim.claim_number, claim.text, cached_report)
                    results[str(cn)] = {
                        "report_md": cached_report,
                        "used_inventions": _used_inventions_for(
                            cached_chain_info, prior_docs, matches=cached_matches
                        ),
                    }
                    continue
            matches, _ = get_matches_from_cache(
                claim,
                prior_docs,
                str(job_dir),
                comparison_mode=compare_mode,
            )
            chain_info = get_claim_chain_info(chain_data, cn) if chain_data else None
            if chain_info and chain_info.get("total"):
                matches, _ = get_matches_from_cache(
                    claim,
                    prior_docs,
                    str(job_dir),
                    allowed_docs=chain_info["total"],
                    comparison_mode=compare_mode,
                )
            secondary_matches = None
            total_refs = chain_info.get("total", []) if chain_info else []
            if len(total_refs) > 1:
                secondary_matches = []
                for sec_idx in total_refs[1:]:
                    sec, _ = get_matches_from_cache(
                        claim,
                        prior_docs,
                        str(job_dir),
                        allowed_docs=[sec_idx],
                        comparison_mode=compare_mode,
                    )
                    secondary_matches.extend(sec)
            batch_items.append((claim, matches, chain_info, secondary_matches))

        if batch_items:
            _update_dependent_batch_status(
                job_id,
                state="running",
                claim_numbers=claim_numbers,
                stage="waiting_for_batch_llm",
                message=f"종속항 {len(batch_items)}개에 대해 LLM 배치 보고서를 생성하고 있습니다.",
                started_at=started_at,
                reports_ready=len(results),
            )
            combined = await _await_with_batch_status_heartbeat(
                generate_dependent_reports_batch(
                    batch_items,
                    prior_docs,
                    settings,
                    prev_context=prev_context if req.use_context else None,
                ),
                job_id=job_id,
                claim_numbers=claim_numbers,
                started_at=started_at,
                stage="waiting_for_batch_llm",
                message_builder=lambda elapsed: (
                    f"종속항 {len(batch_items)}개에 대한 LLM 배치 보고서를 생성 중입니다. ({elapsed})"
                ),
                reports_ready_getter=lambda: len(results),
            )
            combined = sanitize_report_status_icons(_strip_agent_tool_calls(combined))
            parts = _BATCH_SPLIT_RE.split(combined)
            chunks: dict[int, str] = {}
            for i in range(1, len(parts) - 1, 2):
                chunks[int(parts[i])] = parts[i + 1].strip()
            for claim, matches, chain_info, secondary in batch_items:
                raw = chunks.get(claim.claim_number)
                if raw:
                    candidate_report = _assemble_dependent_report(
                        raw,
                        claim,
                        chain_info,
                        settings,
                        matches=matches,
                        secondary_matches=secondary,
                    )
                    if not _has_substantive_dependent_report(candidate_report, claim.claim_number):
                        raw = None
                if not raw:
                    _update_dependent_batch_status(
                        job_id,
                        state="running",
                        claim_numbers=claim_numbers,
                        stage="fallback_single_report",
                        message=f"청구항 {claim.claim_number} 배치 결과가 없어 단건 보고서로 대체합니다.",
                        started_at=started_at,
                        reports_ready=len(results),
                    )
                    raw = await _await_with_batch_status_heartbeat(
                        generate_dependent_report(
                            claim,
                            matches,
                            prior_docs,
                            chain_info,
                            settings,
                            prev_context=prev_context if req.use_context else None,
                            secondary_matches=secondary,
                        ),
                        job_id=job_id,
                        claim_numbers=claim_numbers,
                        started_at=started_at,
                        stage="fallback_single_report",
                        message_builder=lambda elapsed, cn=claim.claim_number: (
                            f"청구항 {cn} 단건 보고서를 생성 중입니다. ({elapsed})"
                        ),
                        reports_ready_getter=lambda: len(results),
                    )
                raw = sanitize_report_status_icons(raw)
                report_md = _assemble_dependent_report(
                    raw,
                    claim,
                    chain_info,
                    settings,
                    matches=matches,
                    secondary_matches=secondary,
                )
                report_md = _prepend_rejection_basis_header(
                    sanitize_report_status_icons(report_md),
                    chain_info,
                    matches,
                )
                if not _has_substantive_dependent_report(report_md, claim.claim_number):
                    raise RuntimeError(
                        f"청구항 {claim.claim_number} 보고서 본문이 비어 있어 저장을 중단했습니다. 다시 생성해 주세요."
                    )
                related_inventions_md = _build_related_inventions_tab(
                    claim,
                    prior_docs,
                    chain_info,
                    job_dir,
                )
                _save_report(job_id, claim.claim_number, report_md)
                _save_reference_db(job_id, claim, matches, prior_docs, chain_info, report_md)
                _save_context_entry(job_id, claim.claim_number, claim.text, report_md)
                results[str(claim.claim_number)] = {
                    "report_md": report_md,
                    "related_inventions_md": related_inventions_md,
                    "used_inventions": _used_inventions_for(
                        chain_info, prior_docs, matches=matches
                    ),
                }
                _update_dependent_batch_status(
                    job_id,
                    state="running",
                    claim_numbers=claim_numbers,
                    stage="saving_reports",
                    message=f"청구항 {claim.claim_number} 보고서가 저장되었습니다.",
                    started_at=started_at,
                    reports_ready=len(results),
                )

        _update_dependent_batch_status(
            job_id,
            state="completed",
            claim_numbers=claim_numbers,
            stage="completed",
            message=f"종속항 보고서 {len(results)}개 생성을 완료했습니다.",
            started_at=started_at,
            reports_ready=len(results),
            completed_at=time.strftime("%Y-%m-%d %H:%M:%S"),
        )
        return {"reports": results}
    except Exception as exc:
        logger.exception("Dependent batch report failed for %s", job_id)
        _update_dependent_batch_status(
            job_id,
            state="failed",
            claim_numbers=claim_numbers,
            stage="failed",
            message="종속항 배치 보고서 생성 중 오류가 발생했습니다.",
            started_at=started_at,
            error=str(exc),
        )
        raise


@router.get("/report_batch_dependent_status/{job_id}")
async def report_batch_dependent_status(job_id: str):
    job_dir = _job_dir(job_id)
    if not job_dir.exists():
        raise HTTPException(status_code=404, detail="작업을 찾을 수 없습니다.")
    return _load_json(_dependent_batch_status_path(job_id), {
        "job_id": job_id,
        "state": "idle",
        "stage": "idle",
        "message": "종속항 배치 보고서 작업을 아직 시작하지 않았습니다.",
        "claim_numbers": [],
        "reports_ready": 0,
        "started_at": "",
        "updated_at": "",
        "completed_at": "",
        "error": "",
    })


@router.post("/chat/{job_id}/{claim_number}")
async def chat_about_report(job_id: str, claim_number: int, req: ChatRequest):
    job_dir = _job_dir(job_id)
    if not job_dir.exists():
        raise HTTPException(status_code=404, detail="작업을 찾을 수 없습니다.")
    if not req.messages:
        raise HTTPException(status_code=400, detail="대화 메시지를 입력해주세요.")

    settings = _load_settings_with_dir()
    claim = next((c for c in _load_claims(job_id) if c.claim_number == claim_number), None)
    claim_text = claim.text if claim else ""
    if req.web_search:
        evidence_rule = (
            "보고서에 있는 인용 결과를 주로 활용하되, 불충분한 경우에는 웹검색으로 보완하세요. "
            "웹검색으로 확인한 문헌은 실제 제목/번호/URL 등을 확인할 수 있는 정보를 포함하세요. "
        )
    else:
        evidence_rule = "보고서에 있는 인용 결과만 활용해 답변하세요. "
    system = (
        "당신은 특허 분석 보조원입니다. "
        f"{evidence_rule}"
        "인용문을 새로 만들지 말고, 근거가 없으면 없다고 답하세요.\n\n"
        f"[청구항 {claim_number}]\n{claim_text}\n\n[보고서]\n{req.report_md}"
    )
    lines = []
    for msg in req.messages[-8:]:
        speaker = "사용자" if msg.role == "user" else "어시스턴트"
        lines.append(f"{speaker}: {msg.content}")
    lines.append("어시스턴트:")
    answer = await call_ai(
        "\n".join(lines),
        system,
        settings,
        agent="compare" if req.web_search else "parser",
        web_search=req.web_search,
    )
    return {"answer": _strip_agent_tool_calls(answer)}


@router.post("/search_missing_prior_art/{job_id}/{claim_number}")
async def search_missing_prior_art(
    job_id: str,
    claim_number: int,
    req: MissingPriorArtSearchRequest,
):
    """미커버 구성만 대상으로 공개 선행기술을 검색한다."""
    job_dir = _job_dir(job_id)
    if not job_dir.exists():
        raise HTTPException(status_code=404, detail="작업을 찾을 수 없습니다.")

    settings = _load_settings_with_dir()
    if settings.engine == "openai":
        raise HTTPException(
            status_code=400,
            detail="현재 OpenAI CLI 실행 경로는 웹 검색 도구를 노출하지 않습니다. Claude 또는 AGY 엔진을 선택해 주세요.",
        )

    claim = next((c for c in _load_claims(job_id) if c.claim_number == claim_number), None)
    if claim is None:
        raise HTTPException(status_code=404, detail=f"청구항 {claim_number}을 찾을 수 없습니다.")

    chain_data = _load_json(job_dir / "citation_chain.json", {})
    assessment = (
        chain_data.get("quantitative_assessment", {}).get(str(claim_number), {})
    )
    requested = {_label.strip().upper() for _label in req.labels if _label.strip()}
    missing = (
        assessment.get("critical_uncovered_labels")
        or assessment.get("uncovered_labels")
        or []
    )
    target_labels = [
        element.label
        for element in claim.elements
        if (not requested and element.label in missing)
        or (requested and element.label.upper() in requested)
    ]
    if not target_labels:
        raise HTTPException(
            status_code=400,
            detail="검색할 미커버 구성이 없습니다. 보고서를 다시 생성하거나 검색할 구성 라벨을 지정해 주세요.",
        )

    targets = [
        {
            "label": element.label,
            "text": element.text,
            "importance": element.importance,
        }
        for element in claim.elements
        if element.label in target_labels
    ]
    chain_info = get_claim_chain_info(chain_data, claim_number) or {}
    used_indices = chain_info.get("total") or []
    prior_docs = _load_prior_docs(job_id)
    used_docs = [
        {
            "index": idx,
            "filename": prior_docs[idx].filename,
            "publication_no": prior_docs[idx].publication_no,
            "title": prior_docs[idx].title,
        }
        for idx in used_indices
        if 0 <= idx < len(prior_docs)
    ]

    search_result = await run_adaptive_prior_art_search(
        claim.text,
        targets,
        used_docs,
        settings,
        additional_query=req.additional_query,
    )
    answer = _strip_agent_tool_calls(search_result.get("result_md", "")).strip()
    if not answer:
        raise HTTPException(status_code=502, detail="웹 검색 결과가 비어 있습니다.")

    payload = {
        "claim_number": claim_number,
        "target_labels": target_labels,
        "targets": targets,
        "used_documents": used_docs,
        "result_md": answer,
        "expanded": search_result.get("expanded", False),
        "search_axes": search_result.get("search_axes", []),
        "initial_stats": search_result.get("initial_stats", {}),
        "final_stats": search_result.get("final_stats", {}),
        "searched_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "engine": settings.engine,
    }
    _write_json(
        _ensure_case_dirs(job_id) / "reports" / f"missing_prior_art_claim{claim_number}.json",
        payload,
    )
    return payload


@router.get("/search_missing_prior_art/{job_id}/{claim_number}")
async def get_missing_prior_art_search(job_id: str, claim_number: int):
    """저장된 미커버 구성 선행기술 검색 결과를 반환한다."""
    result_path = (
        _job_dir(job_id) / "reports" / f"missing_prior_art_claim{claim_number}.json"
    )
    if not result_path.exists():
        raise HTTPException(status_code=404, detail="저장된 선행기술 검색 결과가 없습니다.")
    return _load_json(result_path, {})


@router.post("/cancel")
async def cancel_generation():
    killed = kill_active_cli_procs()
    return {"ok": True, "killed": killed, "message": f"실행 중인 LLM 프로세스 {killed}개를 종료 요청했습니다."}


@router.get("/download/{job_id}/{claim_number}")
async def download_claim(job_id: str, claim_number: int):
    md_path = REPORTS_DIR / f"report_{job_id}_claim{claim_number}.md"
    if not md_path.exists():
        raise HTTPException(status_code=404, detail="보고서가 생성되지 않았습니다.")
    docx_path = _md_to_docx(str(md_path), job_id, claim_number)
    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"report_claim{claim_number}.docx",
    )


@router.get("/download_all/{job_id}")
async def download_all(job_id: str):
    md_files = sorted(REPORTS_DIR.glob(f"report_{job_id}_claim*.md"))
    if not md_files:
        raise HTTPException(status_code=404, detail="생성된 보고서가 없습니다.")
    all_md = "\n\n---\n\n".join(f.read_text(encoding="utf-8") for f in md_files)
    out_md = REPORTS_DIR / f"report_{job_id}_all.md"
    out_md.write_text(all_md, encoding="utf-8")
    docx_path = _md_to_docx_all(all_md, job_id)
    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="report_all.docx",
    )


def _md_to_docx(md_path: str, job_id: str, claim_number: int) -> str:
    from docx import Document
    doc = Document()
    _fill_docx(doc, Path(md_path).read_text(encoding="utf-8"))
    out = REPORTS_DIR / f"report_{job_id}_claim{claim_number}.docx"
    doc.save(str(out))
    return str(out)


def _md_to_docx_all(md_text: str, job_id: str) -> str:
    from docx import Document
    doc = Document()
    _fill_docx(doc, md_text)
    out = REPORTS_DIR / f"report_{job_id}_all.docx"
    doc.save(str(out))
    return str(out)


def _fill_docx(doc, md_text: str) -> None:
    from docx.shared import Pt
    for line in md_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Pt(18)
            if p.runs:
                p.runs[0].italic = True
        elif stripped.startswith("---"):
            doc.add_paragraph("-" * 40)
        else:
            doc.add_paragraph(stripped)


@router.get("/job_status/{job_id}")
async def job_status(job_id: str):
    job_dir = _job_dir(job_id)
    return {
        "exists": job_dir.exists(),
        "prior_count": len(_load_json(job_dir / "prior_docs.json", [])) if job_dir.exists() else 0,
        "claim_count": len(_load_json(job_dir / "claims.json", [])) if job_dir.exists() else 0,
    }


@router.get("/context/{job_id}")
async def get_context(job_id: str):
    context = _load_context(job_id)
    return {"context_claims": context, "count": len(context)}


@router.delete("/context/{job_id}")
async def clear_context(job_id: str):
    job_dir = _job_dir(job_id)
    case_dir = _case_dir(job_id)

    # 사용자가 명시적으로 컨텍스트를 비우면 보고서 연속성과 함께 확정된
    # 인용 체인·문헌 번호 잠금도 초기화한다. 구성대비 캐시는 유지하여 다음
    # 실행에서 같은 원문 근거를 재사용하되, 주·보조문헌과 번호는 새로 선정한다.
    for path in (
        job_dir / "context.json",
        job_dir / "citation_chain.json",
        job_dir / "same_pairs.json",
        case_dir / "parsed" / "citation_chain.json",
    ):
        try:
            path.unlink(missing_ok=True)
        except OSError as exc:
            logger.warning("Could not clear derived context file %s: %s", path, exc)

    report_paths = list(REPORTS_DIR.glob(f"report_{job_id}_claim*.*"))
    report_paths.extend(REPORTS_DIR.glob(f"report_{job_id}_all.*"))
    case_reports_dir = case_dir / "reports"
    if case_reports_dir.exists():
        report_paths.extend(case_reports_dir.glob("claim*.*"))
        report_paths.extend(case_reports_dir.glob("all.*"))
    for path in report_paths:
        try:
            path.unlink(missing_ok=True)
        except OSError as exc:
            logger.warning("Could not clear stale report %s: %s", path, exc)

    for claim in _load_claims(job_id):
        _remove_reference_entries_for_claim(job_id, claim.claim_number)

    return {"ok": True, "citation_numbering_reset": True}


@router.delete("/job/{job_id}")
async def delete_job(job_id: str):
    job_dir = _job_dir(job_id)
    _delete_doc_cache_for_job(job_id)
    if job_dir.exists():
        _rmtree_with_retry(job_dir)
    case_dir = _case_dir(job_id)
    if case_dir.exists():
        _rmtree_with_retry(case_dir)
    for path in REPORTS_DIR.glob(f"report_{job_id}_claim*.*"):
        path.unlink(missing_ok=True)
    for path in REPORTS_DIR.glob(f"report_{job_id}_all.*"):
        path.unlink(missing_ok=True)
    _remove_reference_entries_for_job(job_id)
    return {"ok": True}


@router.delete("/jobs")
async def delete_all_jobs():
    removed = {
        "uploads": 0,
        "cases": 0,
        "reports": 0,
        "doc_cache": 0,
    }

    for path in UPLOADS_DIR.iterdir():
        if path.name.startswith("_"):
            continue
        if path.is_dir():
            _rmtree_with_retry(path)
            removed["uploads"] += 1
        elif path.is_file():
            path.unlink(missing_ok=True)

    for path in CASES_DIR.iterdir():
        if path.is_dir():
            _rmtree_with_retry(path)
            removed["cases"] += 1
        elif path.is_file():
            path.unlink(missing_ok=True)

    for path in REPORTS_DIR.iterdir():
        if path.is_file():
            path.unlink(missing_ok=True)
            removed["reports"] += 1
        elif path.is_dir():
            _rmtree_with_retry(path)
            removed["reports"] += 1

    for path in DOC_CACHE_DIR.iterdir():
        if path.is_file():
            path.unlink(missing_ok=True)
            removed["doc_cache"] += 1
        elif path.is_dir():
            _rmtree_with_retry(path)
            removed["doc_cache"] += 1

    return {"ok": True, "removed": removed}


@router.get("/claim_tree/{job_id}")
async def get_claim_tree(job_id: str):
    job_dir = _job_dir(job_id)
    if not job_dir.exists():
        raise HTTPException(status_code=404, detail="작업을 찾을 수 없습니다.")
    return {
        "job_id": job_id,
        "purpose_effects": _load_json(job_dir / "purpose_effects.json", {
            "purpose": "", "effects": "", "extracted_by": "pending"
        }),
        "claims": _load_json(job_dir / "claims.json", []),
        "same_pairs": _load_json(job_dir / "same_pairs.json", {}),
    }


@router.post("/enhance_purpose/{job_id}")
async def enhance_purpose(job_id: str):
    job_dir = _job_dir(job_id)
    if not job_dir.exists():
        raise HTTPException(status_code=404, detail="작업을 찾을 수 없습니다.")
    settings = _load_settings_with_dir()
    claims = _load_json(job_dir / "claims.json", [])
    independent = [c for c in claims if c.get("claim_type") == "independent"] or claims[:3]
    claims_text = "\n\n".join(f"청구항 {c['claim_number']}:\n{c['text']}" for c in independent)
    result = await enhance_purpose_effects_with_llm(claims_text, settings)
    _write_json(job_dir / "purpose_effects.json", result)
    return result


@router.post("/enhance_claim/{job_id}/{claim_number}")
async def enhance_claim(job_id: str, claim_number: int):
    job_dir = _job_dir(job_id)
    if not job_dir.exists():
        raise HTTPException(status_code=404, detail="작업을 찾을 수 없습니다.")
    settings = _load_settings_with_dir()
    claims = _load_json(job_dir / "claims.json", [])
    claim_data = next((c for c in claims if int(c.get("claim_number", -1)) == claim_number), None)
    if not claim_data:
        raise HTTPException(status_code=404, detail=f"청구항 {claim_number}를 찾을 수 없습니다.")
    enhanced = await enhance_claim_parsing_with_llm(ParsedClaim(**claim_data), settings)
    enhanced_data = enhanced.model_dump()
    updated = [enhanced_data if c.get("claim_number") == claim_number else c for c in claims]
    _write_json(job_dir / "claims.json", updated)
    _write_json(_ensure_case_dirs(job_id) / "parsed" / "claims.json", updated)
    if _claim_analysis_signature(enhanced_data) != _claim_analysis_signature(claim_data):
        _invalidate_claim_derived_artifacts(
            job_id,
            claim_number,
            is_new_claim=False,
            claim_type=enhanced.claim_type,
        )
    return enhanced_data
